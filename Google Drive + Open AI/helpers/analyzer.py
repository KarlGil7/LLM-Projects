import io
import os
import pandas as pd
import zipfile
import chardet
import pdfplumber
from googleapiclient.http import MediaIoBaseDownload
from docx import Document
from openai import OpenAI

# ------------------ OPENAI ------------------
# 🔑 Inicializa OpenAI (usa tu API key)
client = OpenAI()

# ------------------ AUTHENTICATION ------------------
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build

SCOPES = ['https://www.googleapis.com/auth/drive.readonly']

def get_credentials():
    creds = None
    if os.path.exists('../token.json'):
        creds = Credentials.from_authorized_user_file('../token.json', SCOPES)

    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file('../client_secret.json', SCOPES)
            creds = flow.run_local_server(port=0)
        with open('../token.json', 'w') as token:
            token.write(creds.to_json())
    return creds


# ------------------ PARSERS ------------------
def read_docx_from_bytes(raw):
    """Lee todo el texto de un archivo DOCX"""
    doc = Document(io.BytesIO(raw))
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())
    return "\n".join(full_text)

def read_tables_from_docx(raw):
    """Extrae todas las tablas de un DOCX como DataFrames"""
    doc = Document(io.BytesIO(raw))
    tables = []
    for table in doc.tables:
        data = []
        for row in table.rows:
            data.append([cell.text for cell in row.cells])
        tables.append(pd.DataFrame(data))
    return tables

def read_pdf_from_bytes(raw):
    """Lee texto y tablas de un PDF"""
    text_content = []
    tables = []
    with pdfplumber.open(io.BytesIO(raw)) as pdf:
        for page in pdf.pages:
            # texto
            text_content.append(page.extract_text() or "")
            # tablas
            for table in page.extract_tables():
                tables.append(pd.DataFrame(table))
    return {
        "texto": "\n".join(text_content),
        "tablas": tables
    }


def download_file_as_dataframe(service, file_id, mime_type="text/csv"):
    request = service.files().get_media(fileId=file_id)
    fh = io.BytesIO()
    downloader = MediaIoBaseDownload(fh, request)

    done = False
    while not done:
        status, done = downloader.next_chunk()

    fh.seek(0)
    raw = fh.read()

    # 🔎 Caso 1: ZIP (Excel o Word)
    if zipfile.is_zipfile(io.BytesIO(raw)):
        with zipfile.ZipFile(io.BytesIO(raw)) as z:
            if any(name.startswith("xl/") for name in z.namelist()):
                print("📊 Detectado archivo Excel")
                return pd.read_excel(io.BytesIO(raw))
            elif any(name.startswith("word/") for name in z.namelist()):
                print("📄 Detectado archivo Word")
                texto = read_docx_from_bytes(raw)
                tablas = read_tables_from_docx(raw)
                return {"tipo": "word", "texto": texto, "tablas": tablas}
        raise ValueError("❌ ZIP detectado pero no es ni Excel ni Word")

    # 🔎 Caso 2: PDF
    if raw[:4] == b"%PDF":
        print("📑 Detectado archivo PDF")
        return {"tipo": "pdf", **read_pdf_from_bytes(raw)}

    # 🔎 Caso 3: CSV
    enc = chardet.detect(raw)["encoding"] or "utf-8"
    for encoding_try in [enc, "latin1", "cp1252"]:
        try:
            fh2 = io.BytesIO(raw)
            return pd.read_csv(fh2, encoding=encoding_try, engine="python")
        except Exception as e:
            print(f"⚠️ Error leyendo CSV con {encoding_try}: {e}")

    raise ValueError("❌ No se pudo leer el archivo ni como CSV, ni como Excel, ni como Word ni como PDF")


# ------------------ OPENAI QUERIES ------------------
def ask_llm_about_dataframe(df_or_doc, question, context_note="Single file analysis"):
    """
    Envía el contenido del archivo (DataFrame, DOCX o PDF) al LLM de OpenAI.
    """
    if isinstance(df_or_doc, pd.DataFrame):
        # Caso Excel/CSV
        csv_sample = df_or_doc.head(50).to_csv(index=False)
        data_repr = f"CSV/Excel sample (first 50 rows):\n{csv_sample}"

    elif isinstance(df_or_doc, dict):
        # Puede ser Word o PDF
        if "texto" in df_or_doc and "tablas" in df_or_doc:
            texto = df_or_doc.get("texto", "")
            tablas = df_or_doc.get("tablas", [])
            tablas_str = "\n\n".join(
                [t.head(10).to_csv(index=False) for t in tablas]
            ) if tablas else "No tables"

            # Detectamos si es PDF o DOCX solo para etiqueta
            doc_type = "PDF" if df_or_doc.get("source") == "pdf" else "DOCX"
            data_repr = f"""{doc_type} detected.
Text content (first 1000 chars):
{texto[:1000]}...

Tables (first 10 rows each):
{tablas_str}"""

        else:
            raise ValueError("❌ Diccionario detectado pero no tiene formato válido (Word/PDF)")

    else:
        raise ValueError("❌ Tipo de archivo no soportado en ask_llm_about_dataframe")

    prompt = f"""
    You are a data analysis assistant.
    Context: {context_note}
    Here is the document content:

    {data_repr}

    Question: {question}

    Answer clearly and concisely, based only on the provided data.
    """

    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.2
    )

    return response.choices[0].message.content



def compare_two_dataframes(doc1, doc2, question):
    """
    Compara dos documentos que pueden ser:
      - DataFrame (CSV/Excel)
      - dict con {"texto":..., "tablas": [...]} (Word o PDF)
    """
    def summarize_doc(doc, label="Dataset"):
        if isinstance(doc, pd.DataFrame):
            return f"{label} (CSV/Excel, first 30 rows):\n{doc.head(30).to_csv(index=False)}"
        elif isinstance(doc, dict):  # Word o PDF
            texto = doc.get("texto", "")
            tablas = doc.get("tablas", [])
            tablas_str = "\n\n".join(
                [t.head(10).to_csv(index=False) for t in tablas]
            ) if tablas else "No tables"
            return f"""{label} (Document):
Text (first 1000 chars):
{texto[:1000]}...

Tables (first 10 rows each):
{tablas_str}"""
        else:
            return f"{label}: ❌ Unsupported type {type(doc)}"

    doc1_repr = summarize_doc(doc1, "Dataset A")
    doc2_repr = summarize_doc(doc2, "Dataset B")

    prompt = f"""
    You are a comparison assistant.
    Compare the following two datasets/documents:

    {doc1_repr}

    {doc2_repr}

    Task: Compare them to answer this question:
    {question}

    Provide a structured and concise answer.
    """

    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.2
    )

    return response.choices[0].message.content


# ------------------ MAIN ------------------
if __name__ == "__main__":
    creds = get_credentials()
    drive_service = build('drive', 'v3', credentials=creds)

    # ⚡ Ejemplo: analiza un archivo
    # # file_id = "1OcaOReA66xMUkCaXfc6qVfTf3o-ySpCI"  #word
    # file_id = "1ysiqApPoqJvW71jVQfbFroLBl_kJJGdS"  # pdf
    # df = download_file_as_dataframe(drive_service, file_id, mime_type="text/csv")
    # answer = ask_llm_about_dataframe(df, "What is the main content of this document?")
    # print("LLM Answer:\n", answer)

    # ⚡ Ejemplo: compara dos archivos
    file_id1 = "16LRGcPojkC5Q2rEmGwzib5ny1T_Rq-2z"  # PDF
    file_id2 = "14gpCp0scF9DNfGqY0fcFgfOM56d3cofJ"  # PDF
    df1 = download_file_as_dataframe(drive_service, file_id1)
    df2 = download_file_as_dataframe(drive_service, file_id2)
    comparison = compare_two_dataframes(df1, df2, "Compare these files and tell me differences between them")
    print("Comparison Answer:\n", comparison)
